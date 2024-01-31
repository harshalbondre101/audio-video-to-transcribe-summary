import streamlit as st
from openai import OpenAI
from moviepy.editor import *
from docx import Document

model = "gpt-3.5-turbo-16k"

# Streamlit input for the OpenAI API key
api_key = st.text_input("Enter your OpenAI API Key:")
client = OpenAI(api_key=api_key)

# Function to transcribe audio content
def transcribe_audio(audio_content):
    transcript = client.audio.transcriptions.create(
        model="whisper-1",
        file=audio_content
    )
    return transcript

def summarize_transcribe(transcription):
    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": """ You are an advanced AI system specialized in language comprehension, summarization, action item extraction, identification of key points, and sentiment analysis. Below is a piece of text for which I need multiple analyses:

Abstractive Summary:

"Please provide an abstractive summary of the text, condensing it into a concise and coherent paragraph that captures the main points without unnecessary details or tangential information."

Action Item Extraction:

"Identify and list any tasks, assignments, or actions mentioned in the text. These could be specific tasks assigned to individuals or general actions the group intends to take. List these action items clearly and concisely."

Key Points:

"Extract and list the main ideas, findings, or crucial topics discussed in the text. Your goal is to provide a succinct list that highlights the most important aspects of the discussion."

Sentiment Analysis:

"Analyze the sentiment of the text, considering the overall tone, emotions conveyed, and context. Indicate whether the sentiment is positive, negative, or neutral, and provide brief explanations where possible."""

            },
            {
                "role": "user",
                "content": transcription
            }
        ],
        temperature=0
    )
    response_message = response.choices[0].message.content
    return response_message   

# Function to extract audio from video and transcribe it
def process_video_and_transcribe(uploaded_file):
    try:
        with open(uploaded_file.name, "wb") as file_writer:
            file_writer.write(uploaded_file.getbuffer())

        if uploaded_file.type == "video/mp4":
            video = VideoFileClip(uploaded_file.name)
            audio_path = "extracted_audio.wav"
            video.audio.write_audiofile(audio_path, codec='pcm_s16le', fps=44100)
            st.audio(audio_path, format='audio/wav')

            with open(audio_path, 'rb') as audio_file:
                transcript = transcribe_audio(audio_file)
                if transcript:
                    st.subheader("Transcription:")
                    st.text(transcript.text)

                    st.subheader("Summary:")
                    summary = summarize_transcribe(transcript.text)
                    st.text(summary)
                    
                    # Create a .docx document and save the transcription and summary
                    doc = Document()
                    doc.add_heading('Transcription', level=1)
                    doc.add_paragraph(transcript.text)

                    doc.add_heading('Summary', level=1)
                    doc.add_paragraph(summary)

                    doc.save("transcription_summary.docx")
                    st.success("Transcription and Summary saved in 'transcription_summary.docx'")
                    # Add download button for the generated docx file
                    file_name = "transcription_summary.docx"
                    with open(file_name, "rb") as file:
                        btn = st.download_button(
                            label="Download Transcription Summary",
                            data=file,
                            file_name=file_name,
                            mime="application/octet-stream",
                        )
                        
                    if btn:
                        st.success("File downloaded successfully!")

                    
        else:
            st.warning("Please upload a video file.")

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")

# Streamlit app
def main():

    st.title("Audio/Video Transcription App")

    # Upload audio or video file
    uploaded_file = st.file_uploader("Upload a video file (mp4)", type=['mp4'])

    if uploaded_file is not None:
        if uploaded_file.type == 'video/mp4':
            st.video(uploaded_file)
            if st.button("Extract Audio and Transcribe"):
                process_video_and_transcribe(uploaded_file)
        else:
            st.warning("Please upload a video file.")

if __name__ == "__main__":
    main()
